from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r"d:\ZERO-DAY\Section5_Results_Discussion.docx")

FN = "Times New Roman"

def run(para, text, bold=False, italic=False, size=12):
    r = para.add_run(text)
    r.font.name=FN; r.font.size=Pt(size)
    r.font.bold=bold; r.font.italic=italic
    return r

def h(text, size=12, before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(before)
    p.paragraph_format.space_after=Pt(6)
    run(p, text, bold=True, size=size)
    return p

def body(text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(6)
    if indent: p.paragraph_format.first_line_indent=Cm(1.27)
    run(p, text)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(3)
    p.paragraph_format.space_after=Pt(10)
    run(p, text, italic=True, size=10)

def insert_image(path, width_cm=15.0, cap=""):
    try:
        doc.add_picture(path, width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap: caption(cap)
        return True
    except Exception as e:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p, f"[Image not found: {cap}]", italic=True, size=10)
        return False

BASE = r"C:\Users\hp\.gemini\antigravity\brain\d1de8ce3-73c5-4346-bae0-056094385191"

# ── New section: 5.1.6 Trading Screens ────────────────────────────────────────
doc.add_page_break()

h("5.1.6  Live Trading Interface — Candlestick Chart and Order Entry", size=12, before=0)

body(
    "The following screenshots capture the live trading simulation in active operation "
    "during the COV-20 (March 9, 2020) scenario. All prices, news events, and "
    "market conditions shown are historically accurate representations of that day.")

# ── Screenshot A: Coaching Pause Dialog ──────────────────────────────────────
h("A.  In-Session Coaching Pause Dialog", size=11, before=10)
insert_image(
    BASE + r"\main_trading_interface_1777998505569.png", 15.0,
    "Figure 5.6: Coaching Pause Dialog — ORUS pauses the simulation to ask a "
    "contextual multiple-choice question before the user resumes trading"
)
body(
    "At pre-defined points during the simulation, ORUS (the AI coaching system) "
    "pauses the live market clock and presents a Coaching Pause dialog. This "
    "screen tests the user's situational awareness before allowing them to continue. "
    "In the example shown, the question reads: 'NIFTY opened -3% and BRENT is -12%. "
    "What does this most likely mean for the 6 stocks in your watchlist?' Three "
    "answer options are provided. The correct answer (Option A: They will probably "
    "also be red, especially energy and finance) reinforces the concept of index "
    "correlation and sector beta. The background shows the live trading interface "
    "in a dimmed, paused state. A RESUME TRADING button becomes active once the "
    "user selects an answer. This mechanic bridges passive news consumption with "
    "active hypothesis formation, directly supporting the project's first "
    "pedagogical objective.")

# ── Screenshot B: Live Chart + Breaking News Event ───────────────────────────
h("B.  Live Candlestick Chart with Breaking News Event", size=11, before=10)
insert_image(
    BASE + r"\order_entry_sell_active_1777998656575.png", 15.0,
    "Figure 5.7: Live Trading Room — INDIGO 5-minute candlestick chart at 10:21 IST "
    "with Italy lockdown news event popup and Order Ticket at bottom (BUY/SELL, MARKET/LIMIT/SL/SL-M)"
)
body(
    "This screenshot captures the full live trading interface in active operation. "
    "The main panel displays a real-time 5-minute OHLCV candlestick chart for "
    "INDIGO (InterGlobe Aviation Ltd) with overlaid Moving Average lines (MA20 "
    "in white, MA50 in dashed). The current price is Rs. 1,144.44, representing "
    "a gap-down from the previous close of Rs. 1,247.50 (-8.26%). Red candles "
    "dominate the chart, visually encoding the panic sell-off.")
body(
    "A news event popup is displayed in the upper-right area, classified as "
    "CRITICAL · SIGNAL at 09:32 IST: 'Italy extends lockdown to entire country — "
    "60M residents under restriction.' Three stock impact tags are shown below "
    "the headline: INDIGO -1.8% (red, negative impact on aviation), "
    "TITAN -1.0% (red), and SUNPHARMA +1.2% (green, pharma safe-haven rotation). "
    "This teaches users to immediately identify which stocks are affected by a "
    "news catalyst and in which direction.")
body(
    "At the bottom of the screen, the Order Ticket is visible with three sections: "
    "SIDE · SYMBOL (BUY in green, SELL in white toggle), ORDER TYPE (MARKET "
    "selected, with LIMIT, SL, and SL-M as alternatives), and a THESIS field "
    "labelled 'Why this trade? RSI oversold + bounce off support at 1,150...' "
    "which requires the user to write a trading rationale before the order is "
    "accepted. The wallet shows Cash: Rs. 1,00,000 and Portfolio Total: "
    "Rs. 1,00,000. The right rail shows all six stocks with live prices and "
    "percentage changes.")

# ── Screenshot C: Trading Resumes After Circuit Breaker ──────────────────────
h("C.  Post-Circuit-Breaker Resumption — Market Depth Rail", size=11, before=10)
insert_image(
    BASE + r"\market_halt_news_impact_1777998680561.png", 15.0,
    "Figure 5.8: Post-halt resumption at 11:00 IST — 'Trading resumes — NIFTY now "
    "-6.1%' signal event with all six stocks displaying live losses and extended candlestick chart"
)
body(
    "This screenshot captures the simulation at 11:00 IST, after the circuit "
    "breaker halt has expired. A HIGH · SIGNAL news event fires: 'Trading resumes "
    "— NIFTY now -6.1%, all sectoral indices in red.' The candlestick chart for "
    "INDIGO now shows additional price action after the halt, with the capitulation "
    "low clearly visible around the 10:32 mark (where bars become very short, "
    "indicating low price range during the halt).")
body(
    "The right-side watchlist rail shows all six stocks at their live prices: "
    "SUNPHARMA at Rs. 434.13 (+1.38%, green — correctly identified as the pharma "
    "safe-haven), RELIANCE at Rs. 1,239.29 (-8.61%), HDFCBANK at Rs. 1,004.25 "
    "(-4.74%), TITAN at Rs. 1,003.48 (-7.51%), and TCS at Rs. 2,089.43 (-3.11%). "
    "The VWAP (Volume Weighted Average Price) line at Rs. 1,152.15 and the MA20 "
    "at Rs. 1,146.96 are visible on the chart as reference levels for trade entry "
    "decisions. The Order Ticket remains accessible at the bottom for "
    "immediate order placement.")

# ── Screenshot D: Stock Dossier Dialog ────────────────────────────────────────
h("D.  Stock Dossier Dialog — Pre-Trade Chart Analysis", size=11, before=10)
insert_image(
    BASE + r"\stock_dossier_dialog_1777998709231.png", 15.0,
    "Figure 5.9: INDIGO Stock Dossier Modal — Candlestick chart with MA20/MA50, "
    "Support/Resistance levels, Volume bars, and six analysis tabs (Chart, Technicals, "
    "Fundamentals, News, Company, Sector)"
)
body(
    "The Stock Dossier modal opens when a user clicks any stock card in the Prep "
    "Room. This dialog provides comprehensive pre-trade analysis tools across six "
    "tabs: CHART, TECHNICALS, FUNDAMENTALS, NEWS, COMPANY, and SECTOR. "
    "The chart tab (shown) displays INDIGO's 1-month historical price action "
    "rendered as a candlestick chart with three toggle-able overlays: MA20 (solid "
    "yellow/teal), MA50 (dashed), and Support/Resistance levels (S at Rs. 1,219.45 "
    "shown as a green dashed line).")
body(
    "The header shows the ticker symbol INDIGO with company name "
    "InterGlobe Aviation Ltd, sector tag AVIATION · NSE, and the key statistic: "
    "Rs. 1,247.50 with -9.7% over 30 days — flagging significant underperformance "
    "even before the simulation begins. An ORUS TUTOR button in the top-right "
    "opens the AI mentor drawer, where the user can ask ORUS to explain any "
    "pattern or level visible on the chart. A 0/6 STUDIED badge tracks how many "
    "of the six available artifacts (chart, technicals, fundamentals, news, "
    "company, sector) the user has reviewed. The volume histogram at the bottom "
    "shows trading volume per session, with the most recent bars confirming "
    "elevated selling pressure. A CHART TUTORIAL button launches a guided "
    "explanation of how to read the displayed chart patterns.")

out = r"d:\ZERO-DAY\Section5_Results_Discussion_v2.docx"
doc.save(out)
print(f"Saved: {out}")
import os
size = os.path.getsize(out)
print(f"Size: {size/1024/1024:.1f} MB")
