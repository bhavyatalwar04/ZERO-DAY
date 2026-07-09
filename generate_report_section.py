from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(2.54)

# ── Style helpers ─────────────────────────────────────────────────────────────
FONT_NAME   = "Times New Roman"
COLOR_BLACK = RGBColor(0, 0, 0)

def set_run(run, bold=False, italic=False, size=12, color=COLOR_BLACK):
    run.font.name   = FONT_NAME
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color

def heading(doc, text, level, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run(run, bold=True, size=size)
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment  = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_run(run)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    run = p.add_run(text)
    set_run(run)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        set_run(run, bold=True, size=11)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
    # data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(val)
            set_run(run, size=11)
            cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    doc.add_paragraph()   # spacing after table
    return table

def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run(run, italic=True, size=11)

def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(1.27)
        run = p.add_run(line)
        run.font.name   = "Courier New"
        run.font.size   = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "4.  IMPLEMENTATION", level=1, size=14)

body(doc,
    "This chapter presents the complete implementation of ZERO-DAY MARKET, covering "
    "the project structure, the technology stack, and a detailed walkthrough of each "
    "module. The application is built as a full-stack Next.js 16 project using the "
    "App Router paradigm, which co-locates the frontend pages, server-side API routes, "
    "and shared business logic within a single repository. All source code resides in "
    "the frontend/ sub-directory of the repository.", indent=True)

# ── 4.1 Project Structure ─────────────────────────────────────────────────────
heading(doc, "4.1  Project Structure", level=2, size=13)

body(doc,
    "The repository is organised into a top-level directory (ZERO-DAY/) that contains "
    "the Next.js frontend application and a set of experimental Python scripts used for "
    "AI model research. The following directory tree represents the canonical structure "
    "of the project:")

code_block(doc, [
    "ZERO-DAY/",
    "├── frontend/",
    "│   ├── app/                   ← Next.js App Router (pages + API routes)",
    "│   │   ├── page.tsx           ← Landing / Splash page",
    "│   │   ├── layout.tsx         ← Root layout with fonts and providers",
    "│   │   ├── globals.css        ← Global stylesheet",
    "│   │   ├── signup/            ← Authentication page",
    "│   │   ├── welcome/           ← Post-signup onboarding",
    "│   │   ├── onboarding/        ← Knowledge-level selection",
    "│   │   ├── ledger/            ← Scenario library",
    "│   │   ├── sim/[id]/          ← Dynamic simulation routes",
    "│   │   │   ├── prep/          ← Pre-trade briefing room",
    "│   │   │   └── live/          ← Real-time trading simulation",
    "│   │   └── api/               ← Server-side API endpoints",
    "│   │       ├── copilot/       ← AI trading coach (streaming)",
    "│   │       ├── feedback/      ← Post-trade AI feedback",
    "│   │       ├── tutor/         ← Prep Room AI analyst (ORUS)",
    "│   │       ├── sentiment/     ← News headline sentiment scorer",
    "│   │       └── portfolio-feedback/ ← Portfolio run debrief",
    "│   ├── components/            ← Reusable React components",
    "│   │   ├── layout/            ← Sidebar and dashboard shell",
    "│   │   ├── prep/              ← 24 Prep Room components",
    "│   │   ├── live/              ← 11 Live Trading components",
    "│   │   ├── portfolio/         ← 16 Portfolio Game components",
    "│   │   ├── ledger/            ←  9 Scenario Library components",
    "│   │   ├── gameplay/          ← Shared cross-mode components",
    "│   │   └── ui/                ← Design-system primitives",
    "│   ├── lib/                   ← Shared business logic",
    "│   │   ├── data/              ← Static historical data",
    "│   │   │   ├── scenarios.ts   ← 10 scenario definitions",
    "│   │   │   └── scenarios/cov-20/ ← Full intraday data (COV-20)",
    "│   │   ├── contexts/          ← React global state providers",
    "│   │   ├── hooks/             ← Custom game-logic hooks",
    "│   │   ├── utils/             ← Pure utility functions",
    "│   │   └── ai/                ← Groq API client wrapper",
    "│   ├── types/                 ← TypeScript type definitions",
    "│   ├── middleware.ts          ← Supabase auth middleware",
    "│   └── next.config.ts        ← Framework configuration",
    "└── [Python research scripts]  ← AI model experiments",
])

body(doc,
    "The separation of concerns follows the convention established by the Next.js "
    "App Router: pages are co-located with their routes, server-side logic lives "
    "exclusively in app/api/ handlers, and all reusable business logic is extracted "
    "into lib/. This structure ensures that no server-side secrets (API keys) are "
    "ever bundled into the client-side JavaScript.")

# Table 4.1
heading(doc, "Table 4.1 – Top-Level Directory Descriptions", level=3, size=12)
add_table(doc,
    headers=["Directory / File", "Purpose"],
    rows=[
        ["frontend/app/",         "All Next.js pages and server API route handlers"],
        ["frontend/components/",  "Reusable React UI components grouped by feature area"],
        ["frontend/lib/data/",    "Static historical market data (scenarios, OHLCV, news events)"],
        ["frontend/lib/contexts/","React Context providers supplying global application state"],
        ["frontend/lib/hooks/",   "Custom hooks encapsulating game engine logic"],
        ["frontend/lib/utils/",   "Pure utility functions: scoring, maths, audio, telemetry"],
        ["frontend/lib/ai/",      "Groq API client with automatic key-rotation fallback"],
        ["frontend/types/",       "Shared TypeScript interfaces for domain objects"],
        ["middleware.ts",         "Next.js edge middleware enforcing authentication"],
    ],
    col_widths=[5.5, 10.5],
)
caption(doc, "Table 4.1: Description of major project directories")

# ── 4.2 Modules ───────────────────────────────────────────────────────────────
heading(doc, "4.2  Modules", level=2, size=13)

body(doc,
    "The application is decomposed into eight discrete modules. Each module corresponds "
    "to a cohesive concern—authentication, scenario data, game simulation, AI coaching, "
    "scoring, telemetry, audio, and the frontend presentation layer. The sub-sections "
    "below describe each module in detail.")

# 4.2.1
heading(doc, "4.2.1  Authentication Module", level=3, size=12)
body(doc,
    "The authentication module uses Supabase Auth with the @supabase/ssr adapter, "
    "enabling secure, cookie-based sessions that are accessible to both the server "
    "and the browser. The module consists of three components:")
bullet(doc, "Signup / Login page (app/signup/page.tsx) – Collects user credentials and calls "
            "Supabase's signUp / signInWithPassword methods. On success, a session cookie "
            "is written by the SSR adapter.")
bullet(doc, "OAuth callback handler (app/auth/callback/) – Exchanges the authorisation code "
            "returned by Supabase's OAuth flow for a live session.")
bullet(doc, "Edge Middleware (middleware.ts) – Intercepts every incoming HTTP request, "
            "verifies the session token with Supabase, and enforces route protection rules. "
            "Protected routes redirect unauthenticated users to /signup; authenticated "
            "users visiting auth routes are redirected to /dashboard.")

body(doc,
    "A graceful degradation strategy is implemented: if the Supabase environment "
    "variables (NEXT_PUBLIC_SUPABASE_URL and NEXT_PUBLIC_SUPABASE_ANON_KEY) are absent, "
    "the middleware skips all checks and passes requests through unchanged. This allows "
    "the application to run in a local development environment without requiring a "
    "Supabase project.")

# 4.2.2
heading(doc, "4.2.2  Scenario Data Module", level=3, size=12)
body(doc,
    "The scenario data module is the single source of truth for all historical market "
    "events and the price data that drives the simulation. It is implemented as a set "
    "of TypeScript data files under lib/data/.")

body(doc, "Scenario Metadata (lib/data/scenarios.ts)", indent=False)
body(doc,
    "Ten historical scenarios are defined, each conforming to a strongly-typed Scenario "
    "interface. The interface captures the scenario identifier, the ticker symbol, the "
    "historical date and market time, a difficulty rating (1–5), estimated play time, "
    "community solve-rate statistics, a sparkline dataset for visual preview, and a "
    "briefingText field containing pre-formatted HTML-like markup for the dossier overlay.")

add_table(doc,
    headers=["Scenario ID", "Event", "Category", "Difficulty"],
    rows=[
        ["lehman-2008",           "Lehman Brothers Collapse",  "Crash",      "4 – Hard"],
        ["covid-crash-2020",      "COVID Market Panic",        "Crash",      "3 – Medium"],
        ["gamestop-2021",         "GameStop Short Squeeze",    "Social",     "5 – Expert"],
        ["apple-earnings-2020",   "Apple Surprise Quarter",    "Earnings",   "2 – Easy"],
        ["tesla-earnings-2019",   "Tesla Shocks the Bears",    "Earnings",   "3 – Medium"],
        ["netflix-miss-2019",     "Netflix Subscriber Shock",  "Earnings",   "3 – Medium"],
        ["brexit-2016",           "Brexit Referendum Shock",   "Economic",   "4 – Hard"],
        ["flash-crash-2010",      "The Flash Crash of 2010",   "Technical",  "4 – Hard"],
        ["crypto-winter-2018",    "Crypto Winter 2018",        "Black Swan", "5 – Expert"],
        ["india-demonetization",  "India Demonetization 2016", "Economic",   "3 – Medium"],
    ],
    col_widths=[4.5, 5.5, 3.5, 2.5],
)
caption(doc, "Table 4.2: Historical scenarios available in the Scenario Library")

body(doc, "Intraday Price Data — COV-20 (lib/data/scenarios/cov-20/)", indent=False)
body(doc,
    "The COVID-19 crash scenario (March 9, 2020) is the first fully-implemented live "
    "scenario, comprising six sub-files:")
bullet(doc, "timeline.ts – Generates 75 five-minute OHLCV (Open-High-Low-Close-Volume) "
            "bars for six Indian stocks (INDIGO, SUNPHARMA, RELIANCE, HDFCBANK, TITAN, TCS) "
            "using a deterministic pseudo-random number generator (mulberry32). Historical "
            "targets for the open percentage, intraday low percentage, and closing percentage "
            "are encoded for each stock; the PRNG adds realistic micro-noise while keeping "
            "the chart identical across page reloads.")
bullet(doc, "live-events.ts – Defines twenty time-stamped news events, eleven ORUS coaching "
            "whispers, and one circuit-breaker halt event. Each news event carries a severity "
            "rating and a classification of either signal (actionable market information) or "
            "noise (distracting but irrelevant information), training users to distinguish "
            "between the two.")
bullet(doc, "companies.ts and stocks.ts – Static fundamental data used in the Prep Room "
            "dossier cards.")
bullet(doc, "macro.ts – Macro-economic backdrop data (global indices, commodity prices) "
            "displayed in the macro-intelligence rail.")

# 4.2.3
heading(doc, "4.2.3  Game Engine Module", level=3, size=12)
body(doc,
    "The game engine module contains all real-time simulation logic. It is split into "
    "two engines serving the two game modes.")

body(doc, "A.  Live Trading Engine (lib/contexts/live-session-context.tsx)", indent=False)
body(doc,
    "The live trading engine is implemented as a React Context backed by a useReducer "
    "state machine. The full session state is represented as a single immutable object "
    "updated exclusively through dispatched actions, ensuring predictable state transitions.")

add_table(doc,
    headers=["Action Type", "Effect"],
    rows=[
        ["START",        "Transitions status from PRE_OPEN to LIVE"],
        ["TICK",         "Advances currentMinute, fires circuit breakers, matches orders, records equity point"],
        ["PLACE_ORDER",  "Validates funds/position, appends order; MARKET orders fill immediately"],
        ["CANCEL_ORDER", "Sets order status to CANCELLED"],
        ["SET_STOP",     "Attaches a stop-loss price to an open position"],
        ["PAUSE / RESUME","Toggles session between LIVE and PAUSED"],
        ["SKIP_HALT",    "Jumps time to end of circuit-breaker halt"],
        ["END",          "Closes the session; triggers auto square-off"],
        ["SET_SPEED",    "Changes tick rate: 1× (1500 ms), 5× (300 ms), 10× (150 ms)"],
    ],
    col_widths=[4.0, 12.0],
)
caption(doc, "Table 4.3: Live Trading Engine reducer actions")

body(doc,
    "The tick loop fires at an interval determined by the selected speed multiplier. "
    "On every tick, the reducer: (1) advances the simulated clock; (2) checks for "
    "circuit-breaker triggers and halts the session if required; (3) calls matchOrders() "
    "to fill any pending LIMIT, SL, or SL-M orders whose trigger price has been "
    "crossed; and (4) appends a data point to the running equity curve. At session end "
    "(minute 375, representing 15:30 IST), all open positions are automatically squared "
    "off at the prevailing market price.")

body(doc, "B.  Portfolio Construction Engine (lib/hooks/use-portfolio-game.ts)", indent=False)
body(doc,
    "The portfolio game engine is implemented as a custom React hook (usePortfolioGame) "
    "that manages an eight-state machine: intro → allocating → running → rebalance → "
    "flash_crash → running → closed → results. Users allocate virtual capital "
    "(₹1,00,000) across six stocks at the start. During the running phase, four scripted "
    "news events pause the game and open fifteen-second rebalance windows, forcing the "
    "user to make portfolio adjustment decisions under time pressure. A flash-crash "
    "event presents a binary decision: panic-sell all positions at distressed prices, "
    "or hold through the volatility. At game end, the engine computes:")
bullet(doc, "Final P&L in rupees and percentage")
bullet(doc, "Diversification score using the Herfindahl-Hirschman Index (HHI), scaled 0–5")
bullet(doc, "A do-nothing benchmark (performance if initial allocation was held unchanged)")
bullet(doc, "A perfect-play benchmark (performance of the optimal sector-rotation strategy)")
bullet(doc, "A list of detected mistakes (panic sell, over-concentration, wrong sector hold)")
bullet(doc, "Causal rules unlocked by the user's observed behaviour")

# 4.2.4
heading(doc, "4.2.4  Artificial Intelligence Module", level=3, size=12)
body(doc,
    "The AI module provides five server-side API endpoints, all powered by the Groq "
    "inference API running the llama-3.1-8b-instant large language model. Groq is "
    "selected for its sub-500 ms response latency, which is essential for a real-time "
    "trading environment. A key-rotation fallback wrapper (lib/ai/groq-client.ts) "
    "distributes requests across up to four API keys; if a key returns HTTP 429 "
    "(rate-limited), the next key is tried automatically.")

add_table(doc,
    headers=["Endpoint", "Mode", "Persona / Purpose", "Output Format"],
    rows=[
        ["/api/copilot",            "Streaming",     "Expert quant teacher; answers live trading questions with injected market context (price, RSI, MACD, current news)", "Plain text stream"],
        ["/api/copilot (portfolio)", "Non-streaming", "Portfolio whisper coach; names cognitive biases in ≤20 words",  "Plain text"],
        ["/api/feedback",           "Streaming",     "Grades the user's prediction; returns assessment, insight, and a historical trading quote", "Streaming JSON"],
        ["/api/tutor",              "Non-streaming", "ORUS — a clinical, sarcastic mentor; applies a trading concept to specific visible data",   "Plain text"],
        ["/api/sentiment",          "Non-streaming", "Quantitative NLP model; scores a news headline as bullish / bearish / neutral (0–100)",     "JSON {score, type}"],
        ["/api/portfolio-feedback", "Non-streaming", "Trading coach; cites specific ₹ decisions and names one cognitive bias",                   "JSON {headline, winMoment, costMoment, bias, proTip}"],
    ],
    col_widths=[4.0, 2.5, 6.5, 3.0],
)
caption(doc, "Table 4.4: AI API endpoints, modes, and output formats")

body(doc,
    "Each endpoint constructs a carefully engineered system prompt that constrains "
    "the model's persona, output length, and format. For the portfolio feedback "
    "endpoint, a buildSummary() function first converts the PortfolioRunResult object "
    "into a plain-English narrative (allocation snapshots, detected mistakes, "
    "diversification score, rules unlocked) before passing it to the model, ensuring "
    "the AI feedback is grounded in the user's actual decisions.")

# 4.2.5
heading(doc, "4.2.5  Scoring and Progression Module", level=3, size=12)
body(doc,
    "The scoring and progression module persists all user statistics to the browser's "
    "localStorage under the key zdm_user_v2. The stored structure tracks XP, level, "
    "daily streak, completed scenarios, accuracy history, and unlocked badges.")

body(doc, "XP Calculation (lib/utils/scoring.ts)", indent=False)
body(doc,
    "Each scenario prediction is scored across three dimensions:")
bullet(doc, "Direction Score: 50 points if the predicted price direction (up / down / flat) matches the historical outcome, otherwise 0.")
bullet(doc, "Magnitude Score: Up to 50 points based on proximity of the predicted percentage magnitude to the actual outcome. A perfect magnitude prediction earns 50 points; the score decays linearly.")
bullet(doc, "Confidence Bonus: A multiplier based on the user's stated confidence (1–5 stars). High confidence on a correct call earns +20 points; high confidence on a wrong call incurs a −10 penalty.")
bullet(doc, "Total XP = max(0, direction + magnitude + confidence) × 1.5")

body(doc, "Levelling System", indent=False)
body(doc,
    "Seven XP thresholds define the level ladder: 0, 500, 1500, 3000, 6000, 10,000, "
    "and 20,000 XP. Separately, a rank title (Novice Trader → Master Trader) is awarded "
    "based on the number of completed scenarios. Five badge types are checked after "
    "every session: first-trade, first-correct, streak-7, 10-scenarios, and crash-direction.")

# 4.2.6
heading(doc, "4.2.6  Telemetry Module", level=3, size=12)
body(doc,
    "The telemetry module (lib/utils/telemetry.ts) silently records user behaviour "
    "during the Prep Room session to enable data-driven post-game coaching. Events are "
    "time-stamped relative to the session start and persisted to localStorage under "
    "zdm_prep_telemetry_{scenarioId}. The following event types are captured:")

add_table(doc,
    headers=["Event Type", "Data Captured"],
    rows=[
        ["artifact-view",         "Stock symbol, artifact name, time spent viewing (ms)"],
        ["tutor-open",            "Stock symbol and artifact that triggered the AI query"],
        ["tutor-marked-studied",  "Confirmation that user acknowledged the AI explanation"],
        ["tab-blur",              "Timestamp when user switched away from the application"],
        ["allocation-change",     "Stock symbol and new rupee allocation value"],
        ["deploy-clicked",        "Elapsed seconds from session start to deploy action"],
    ],
    col_widths=[4.5, 11.5],
)
caption(doc, "Table 4.5: Telemetry event types recorded during the Prep Room session")

body(doc,
    "The summarizeSession() function aggregates the raw event log into derived metrics: "
    "total session duration, time spent per artifact, tutor open count, list of "
    "artifacts skipped before deploying, number of allocation changes, and time-to-deploy. "
    "These metrics feed the post-session ORUS debrief.")

# 4.2.7
heading(doc, "4.2.7  Audio Module", level=3, size=12)
body(doc,
    "The audio module (lib/utils/sound.ts) synthesises all in-game sound effects using "
    "the browser's Web Audio API with no external library dependency. Eight named sounds "
    "are defined: opening-bell, closing-bell, breaking-news, flash-crash, tick-soft, "
    "rule-unlock, buy, and sell. Each sound is synthesised by scheduling one or more "
    "oscillator nodes (sine, triangle, square, or sawtooth waveform) with gain envelopes "
    "that produce realistic attack and decay characteristics. A throttle mechanism "
    "prevents high-frequency sounds (such as tick-soft) from firing more than once every "
    "four seconds. A global mute flag allows users to silence all audio without "
    "disrupting the game state.")

# 4.2.8
heading(doc, "4.2.8  Frontend Presentation Module", level=3, size=12)
body(doc,
    "The frontend presentation module encompasses all React components and the design "
    "system that governs the visual language of the application.")

body(doc, "Design System", indent=False)
body(doc,
    "The application uses a cinematic terminal aesthetic inspired by professional "
    "trading terminals and espionage film interfaces. The colour palette is built on a "
    "near-black background (#050505), a primary accent of red (#dc2626), and off-white "
    "text (#e8e8e8). Monetary gains are rendered in green (#22c55e) and losses in "
    "red (#ef4444). Typography uses Anton for large display headlines, Geist Sans for "
    "interface labels, and Geist Mono for all numeric values with tabular-nums enabled "
    "to prevent layout shift during live price updates. Animation is handled by Framer "
    "Motion throughout, with staggered entry transitions on all page loads.")

body(doc, "Component Organisation", indent=False)
body(doc,
    "Components are grouped into seven directories, each scoped to a specific feature "
    "area. Shared design-system primitives (Button, Card, Badge, Input, Logo) live in "
    "components/ui/ and enforce consistent styling. Feature-specific components are "
    "never imported across feature boundaries; all cross-feature sharing passes through "
    "the ui/ primitives or the lib/ utilities.")

add_table(doc,
    headers=["Component Directory", "Count", "Key Components"],
    rows=[
        ["components/prep/",      "24", "DossierWorkbench, StockCardGrid, Artifacts, TutorDrawer, MacroIntelRail, AllocPanel, PrepHUD"],
        ["components/live/",      "11", "LiveChart, LiveHUD, BottomDock, RightRail, Overlays, LiveTutorial, LiveCoachPrompts"],
        ["components/portfolio/", "16", "TradingFloor, StockCard, RebalanceModal, FlashCrashOverlay, CausalPulseOverlay, Results"],
        ["components/ledger/",    " 9", "BookFrame, BriefingOverlay, BookmarkShelf, LedgerHUD, MarketChronicleTicker"],
        ["components/gameplay/",  " 6", "AiTutor, CandlestickChart, MarketDrawer, NewsDrawer, ScenarioChart"],
        ["components/layout/",    " 2", "Sidebar, DashboardContent"],
        ["components/ui/",        "11", "Button, Card, Badge, Input, Logo, ParticleTextEffect, SpookySmokeAnimation"],
    ],
    col_widths=[4.5, 1.5, 10.0],
)
caption(doc, "Table 4.6: Component directories, counts, and representative components")

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"d:\ZERO-DAY\Section4_Implementation.docx"
doc.save(out)
print(f"Saved: {out}")
